import logging
from docx import Document
from app.parsers.base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """DOCX 파일 파서"""

    async def parse(self, file_path: str) -> ParseResult:
        """DOCX에서 텍스트 추출"""
        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"DOCX 파일을 열 수 없습니다: {e}")

        text_content = []

        # 문단 추출
        try:
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
        except Exception as e:
            logger.warning(f"DOCX 문단 추출 중 오류: {e}")

        # 표 추출 (일정표 추출에 중요)
        try:
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text.strip():
                    text_content.append(f"\n[표]\n{table_text}")
        except Exception as e:
            logger.warning(f"DOCX 표 추출 중 오류: {e}")

        return ParseResult(text="\n\n".join(text_content))

    def _extract_table(self, table) -> str:
        """표에서 텍스트 추출"""
        result = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # 중복 셀 제거 (병합된 셀 처리)
            unique_cells = []
            prev = None
            for cell in cells:
                if cell != prev:
                    unique_cells.append(cell)
                    prev = cell
            result.append(" | ".join(unique_cells))
        return "\n".join(result)
